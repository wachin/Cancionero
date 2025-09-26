#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Unificador de archivos (TXT / DOCX / ODT) con PyQt6

Cambios solicitados:
• Tres botones de añadir: «Añadir ODT», «Añadir DOCX», «Añadir TXT»
• Eliminado el selector «Modo de unión …»
• Arrastrar y soltar (Drag & Drop) implementado en la ventana principal (QMainWindow)
• Unir TXT preservando contenido
• Unir DOCX preservando formato (docxcompose)
• Unir ODT preservando formato vía LibreOffice headless (ODT→DOCX→ODT opcional)
• Productividad: ordenar, quitar duplicados, validar tipos, mover arriba/abajo,
  seleccionar todo/invertir, limpiar lista, abrir carpeta de salida
• Internacionalización: carga de traducciones del sistema (qt6-translations-l10n)

Requisitos Debian 12 (sugeridos):
    sudo apt-get update
    sudo apt-get install -y python3-pyqt6 python3-docx python3-docxcompose libreoffice qt6-translations-l10n
    # (opcional) unoconv

NOTA: Para ODT es necesario tener 'soffice' (LibreOffice) en PATH.
"""

from __future__ import annotations
import sys
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from PyQt6.QtCore import Qt, QLocale, QLibraryInfo, QTranslator
from PyQt6.QtGui import QIcon, QDragEnterEvent, QDropEvent
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QListWidget, QListWidgetItem,
    QPushButton, QFileDialog, QLabel, QLineEdit, QMessageBox, QFrame
)

# ------------------------------- Utilidades ------------------------------- #

def ensure_parent_dir(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)

def which(cmd: str) -> str | None:
    for p in os.environ.get("PATH", "").split(os.pathsep):
        cand = Path(p) / cmd
        if cand.exists() and os.access(cand, os.X_OK):
            return str(cand)
    return None

VALID_EXTS = {'.txt', '.docx', '.odt'}

# ------------------------------ Merge engines ----------------------------- #

def merge_txt(files: list[Path], out_path: Path) -> None:
    ensure_parent_dir(out_path)
    with out_path.open('w', encoding='utf-8') as fout:
        for i, p in enumerate(files):
            try:
                content = p.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                content = p.read_text(encoding='latin-1')
            if i != 0:
                fout.write('

---- FIN DE ARCHIVO ----

')
            fout.write(f"<!-- Inicio: {p.name} -->
")
            fout.write(content)
            fout.write(f"
<!-- Fin: {p.name} -->
")


def merge_docx(files: list[Path], out_path: Path) -> None:
    try:
        from docx import Document
        from docxcompose.composer import Composer
    except Exception as e:
        raise RuntimeError("Faltan dependencias: python3-docx y python3-docxcompose") from e

    ensure_parent_dir(out_path)
    base = Document(str(files[0]))
    composer = Composer(base)
    for f in files[1:]:
        composer.append(Document(str(f)))
    composer.save(str(out_path))


def soffice_convert(input_path: Path, to_ext: str, out_dir: Path) -> Path:
    soffice = which('soffice')
    if not soffice:
        raise RuntimeError("No se encontró 'soffice'. Instala LibreOffice: sudo apt-get install -y libreoffice")
    cmd = [soffice, '--headless', '--convert-to', to_ext, '--outdir', str(out_dir), str(input_path)]
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if proc.returncode != 0:
        raise RuntimeError(f"Error convirtiendo {input_path.name} a {to_ext}: {proc.stderr.decode(errors='ignore')}")
    produced = out_dir / f"{input_path.stem}.{to_ext}"
    if not produced.exists():
        candidates = list(out_dir.glob(f"{input_path.stem}*.{to_ext}"))
        if not candidates:
            raise RuntimeError(f"Conversión no generó salida para {input_path.name}")
        produced = candidates[0]
    return produced


def merge_odt(files: list[Path], out_path: Path) -> None:
    try:
        from docx import Document
        from docxcompose.composer import Composer
    except Exception as e:
        raise RuntimeError("Faltan dependencias: python3-docx y python3-docxcompose") from e

    ensure_parent_dir(out_path)
    with tempfile.TemporaryDirectory() as td:
        tmp = Path(td)
        docx_list: list[Path] = []
        for f in files:
            if f.suffix.lower() != '.odt':
                raise ValueError(f"Archivo no ODT en unión ODT: {f}")
            converted = soffice_convert(f, 'docx', tmp)
            docx_list.append(converted)

        merged_docx = tmp / 'merged.docx'
        base = Document(str(docx_list[0]))
        composer = Composer(base)
        for p in docx_list[1:]:
            composer.append(Document(str(p)))
        composer.save(str(merged_docx))

        if out_path.suffix.lower() == '.odt':
            produced = soffice_convert(merged_docx, 'odt', out_path.parent)
            if produced != out_path:
                if out_path.exists():
                    out_path.unlink()
                produced.rename(out_path)
        else:
            shutil.copy2(merged_docx, out_path)

# ------------------------------ Ventana UI -------------------------------- #

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Unir archivos (TXT / DOCX / ODT)')
        self.resize(950, 620)
        self.setAcceptDrops(True)  # ✅ aceptar D&D en la ventana principal

        self.files: list[Path] = []
        self.output_folder: Path | None = None

        # --- Contenido central ---
        central = QWidget(self)
        self.setCentralWidget(central)
        main = QVBoxLayout(central)

        # Lista
        self.list_widget = QListWidget()
        self.list_widget.setFrameShape(QFrame.Shape.Box)
        self.list_widget.setAlternatingRowColors(True)
        self.list_widget.setSelectionMode(self.list_widget.SelectionMode.ExtendedSelection)
        self.list_widget.setAcceptDrops(False)  # ✅ que no capture el drop: lo maneja la ventana
        main.addWidget(self.list_widget, 1)

        # Botonera de añadir (tres botones)
        row_add = QHBoxLayout()
        btn_add_odt = QPushButton('Añadir ODT')
        btn_add_odt.clicked.connect(lambda: self.pick_and_add(['*.odt']))
        row_add.addWidget(btn_add_odt)

        btn_add_docx = QPushButton('Añadir DOCX')
        btn_add_docx.clicked.connect(lambda: self.pick_and_add(['*.docx']))
        row_add.addWidget(btn_add_docx)

        btn_add_txt = QPushButton('Añadir TXT')
        btn_add_txt.clicked.connect(lambda: self.pick_and_add(['*.txt']))
        row_add.addWidget(btn_add_txt)
        main.addLayout(row_add)

        # Botonera de lista
        row = QHBoxLayout()
        btn_remove = QPushButton('Eliminar selección')
        btn_remove.clicked.connect(self.on_remove_selected)
        row.addWidget(btn_remove)

        btn_clear = QPushButton('Vaciar lista')
        btn_clear.clicked.connect(self.on_clear)
        row.addWidget(btn_clear)

        btn_up = QPushButton('Mover arriba')
        btn_up.clicked.connect(self.on_move_up)
        row.addWidget(btn_up)

        btn_down = QPushButton('Mover abajo')
        btn_down.clicked.connect(self.on_move_down)
        row.addWidget(btn_down)
        main.addLayout(row)

        # Productividad
        prod = QHBoxLayout()
        btn_sort = QPushButton('Ordenar A–Z')
        btn_sort.clicked.connect(self.on_sort)
        prod.addWidget(btn_sort)

        btn_dedup = QPushButton('Quitar duplicados')
        btn_dedup.clicked.connect(self.on_dedup)
        prod.addWidget(btn_dedup)

        btn_validate = QPushButton('Validar tipos contra salida')
        btn_validate.clicked.connect(self.on_validate_types)
        prod.addWidget(btn_validate)

        btn_select_all = QPushButton('Seleccionar todo')
        btn_select_all.clicked.connect(self.on_select_all)
        prod.addWidget(btn_select_all)

        btn_invert = QPushButton('Invertir selección')
        btn_invert.clicked.connect(self.on_invert_selection)
        prod.addWidget(btn_invert)
        main.addLayout(prod)

        # Salida
        out_row1 = QHBoxLayout()
        out_row1.addWidget(QLabel('Nombre de salida (la extensión define el modo):'))
        self.output_name = QLineEdit('merged.docx')  # por defecto DOCX
        out_row1.addWidget(self.output_name, 1)
        main.addLayout(out_row1)

        out_row2 = QHBoxLayout()
        self.output_folder_label = QLabel('Carpeta de salida: (automática)')
        out_row2.addWidget(self.output_folder_label, 1)
        btn_pick_folder = QPushButton('Elegir carpeta…')
        btn_pick_folder.clicked.connect(self.on_pick_folder)
        out_row2.addWidget(btn_pick_folder)
        btn_open_folder = QPushButton('Abrir carpeta')
        btn_open_folder.clicked.connect(self.on_open_folder)
        out_row2.addWidget(btn_open_folder)
        main.addLayout(out_row2)

        # Ejecutar
        run_row = QHBoxLayout()
        btn_join = QPushButton('Unir')
        btn_join.clicked.connect(self.on_join)
        run_row.addStretch(1)
        run_row.addWidget(btn_join)
        main.addLayout(run_row)

    # ----------------------- Drag & Drop en ventana ----------------------- #
    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event: QDropEvent):
        if not event.mimeData().hasUrls():
            event.ignore()
            return
        paths = []
        for url in event.mimeData().urls():
            if url.isLocalFile():
                p = Path(url.toLocalFile())
                if p.exists() and p.is_file() and p.suffix.lower() in VALID_EXTS:
                    paths.append(p)
        self.add_files(paths)
        event.acceptProposedAction()

    # -------------------------- Gestión de lista ------------------------- #
    def pick_and_add(self, patterns: list[str]):
        filt = ';;'.join(f'Tipo ({pat})' for pat in patterns) + ';;Todos (*)'
        files, _ = QFileDialog.getOpenFileNames(self, 'Seleccionar archivos', str(Path.home()), filt)
        paths = [Path(f) for f in files]
        # Filtra por patrones
        allowed = {p.split('*')[-1].lower() for p in patterns}  # {'.odt', '.docx', '.txt'}
        paths = [p for p in paths if p.suffix.lower() in allowed]
        self.add_files(paths)

    def add_files(self, paths: list[Path]):
        if not paths:
            return
        added = False
        for p in paths:
            p = Path(p)
            if p.suffix.lower() not in VALID_EXTS:
                continue
            if p not in self.files:
                self.files.append(p)
                self.list_widget.addItem(QListWidgetItem(str(p)))
                added = True
        if added and self.output_folder is None and self.files:
            self.set_output_folder(self.files[0].parent)

    def on_remove_selected(self):
        rows = sorted({self.list_widget.row(it) for it in self.list_widget.selectedItems()}, reverse=True)
        for r in rows:
            self.list_widget.takeItem(r)
            if 0 <= r < len(self.files):
                del self.files[r]
        if not self.files:
            self.set_output_folder(None)

    def on_clear(self):
        self.list_widget.clear()
        self.files.clear()
        self.set_output_folder(None)

    def on_move_up(self):
        items = self.list_widget.selectedItems()
        if not items:
            return
        rows = sorted({self.list_widget.row(it) for it in items})
        for r in rows:
            if r == 0:
                continue
            self.files[r-1], self.files[r] = self.files[r], self.files[r-1]
            it = self.list_widget.takeItem(r)
            self.list_widget.insertItem(r-1, it)
            it.setSelected(True)

    def on_move_down(self):
        items = self.list_widget.selectedItems()
        if not items:
            return
        rows = sorted({self.list_widget.row(it) for it in items}, reverse=True)
        for r in rows:
            if r >= self.list_widget.count()-1:
                continue
            self.files[r+1], self.files[r] = self.files[r], self.files[r+1]
            it = self.list_widget.takeItem(r)
            self.list_widget.insertItem(r+1, it)
            it.setSelected(True)

    def on_sort(self):
        paired = list(zip(self.files, [self.list_widget.item(i).text() for i in range(self.list_widget.count())]))
        paired.sort(key=lambda x: (x[0].suffix.lower(), x[0].name.lower()))
        self.files = [p for p, _ in paired]
        self.list_widget.clear()
        for p in self.files:
            self.list_widget.addItem(QListWidgetItem(str(p)))

    def on_dedup(self):
        seen = set()
        new_files = []
        for p in self.files:
            if p not in seen:
                seen.add(p)
                new_files.append(p)
        if len(new_files) != len(self.files):
            self.files = new_files
            self.list_widget.clear()
            for p in self.files:
                self.list_widget.addItem(QListWidgetItem(str(p)))

    def on_validate_types(self):
        # Determina modo a partir de la extensión del archivo de salida
        out_ext = Path(self.output_name.text().strip() or 'merged.docx').suffix.lower()
        if out_ext not in {'.txt', '.docx', '.odt'}:
            QMessageBox.warning(self, 'Extensión no válida', 'El nombre de salida debe terminar en .txt, .docx o .odt')
            return
        bad = [p for p in self.files if p.suffix.lower() != out_ext]
        if bad:
            QMessageBox.warning(self, 'Tipos no válidos', 'Estos archivos no coinciden con la extensión de salida ' + out_ext + ' :
' + '
'.join(str(p) for p in bad))
        else:
            QMessageBox.information(self, 'OK', 'Todos los archivos coinciden con el tipo de salida.')

    def on_select_all(self):
        self.list_widget.selectAll()

    def on_invert_selection(self):
        for i in range(self.list_widget.count()):
            it = self.list_widget.item(i)
            it.setSelected(not it.isSelected())

    # ----------------------- Carpeta/nombre de salida --------------------- #
    def set_output_folder(self, folder: Path | None):
        self.output_folder = folder
        if folder is None:
            self.output_folder_label.setText('Carpeta de salida: (automática)')
        else:
            self.output_folder_label.setText(f'Carpeta de salida: {folder}')

    def on_pick_folder(self):
        d = QFileDialog.getExistingDirectory(self, 'Elegir carpeta de salida', str(self.output_folder or Path.home()))
        if d:
            self.set_output_folder(Path(d))

    def on_open_folder(self):
        folder = self.output_folder or (self.files[0].parent if self.files else Path.home())
        if sys.platform.startswith('linux'):
            subprocess.Popen(['xdg-open', str(folder)])
        elif sys.platform == 'darwin':
            subprocess.Popen(['open', str(folder)])
        elif os.name == 'nt':
            os.startfile(str(folder))  # type: ignore

    # ----------------------------- Unir ---------------------------------- #
    def on_join(self):
        if not self.files:
            QMessageBox.warning(self, 'Sin archivos', 'Añade archivos a la lista.')
            return
        name = self.output_name.text().strip() or 'merged.docx'
        out_ext = Path(name).suffix.lower()
        if out_ext not in {'.txt', '.docx', '.odt'}:
            QMessageBox.warning(self, 'Extensión no válida', 'El nombre de salida debe terminar en .txt, .docx o .odt')
            return
        if not name.lower().endswith(out_ext):
            name += out_ext
        out_folder = self.output_folder or self.files[0].parent
        out_path = (out_folder / name).resolve()

        # Validar tipos contra la extensión de salida
        wrong = [p for p in self.files if p.suffix.lower() != out_ext]
        if wrong:
            QMessageBox.warning(self, 'Tipos mezclados', 'Estos archivos no coinciden con la extensión de salida ' + out_ext + ' :
' + '
'.join(str(p) for p in wrong))
            return

        try:
            if out_ext == '.txt':
                merge_txt(self.files, out_path)
            elif out_ext == '.docx':
                merge_docx(self.files, out_path)
            else:  # .odt
                merge_odt(self.files, out_path)
            QMessageBox.information(self, 'Completado', f'Se creó el archivo:
{out_path}')
        except Exception as e:
            QMessageBox.critical(self, 'Error', str(e))

# ----------------------------- i18n Qt6 ----------------------------------- #

def install_translators(app: QApplication):
    """Carga traducciones del sistema si están instaladas (qt6-translations-l10n)."""
    try:
        locale = QLocale.system()
        qt_translator = QTranslator(app)
        qtbase_translator = QTranslator(app)
        qt_path = QLibraryInfo.path(QLibraryInfo.LibraryPath.TranslationsPath)
        qt_translator.load(locale, 'qt', '_', qt_path)
        qtbase_translator.load(locale, 'qtbase', '_', qt_path)
        app.installTranslator(qt_translator)
        app.installTranslator(qtbase_translator)
    except Exception:
        pass

# --------------------------------- main ----------------------------------- #

def main():
    app = QApplication(sys.argv)
    install_translators(app)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
